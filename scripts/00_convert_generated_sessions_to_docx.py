# convert_answer_to_docx.py
import argparse
import json
import yaml
from pathlib import Path

# Third-party libraries: pip install python-docx pyyaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# The create_session_docx function is correct and does not need changes.
# I've included it here for completeness.
def create_session_docx(session_data: dict, output_dir: Path):
    """Creates and saves a single .docx file for a given squash session."""
    try:
        session_id = session_data.get("session_id", "unknown_session")
        question = session_data.get("query_text", "N/A")
        text_content = session_data.get("generated_plan", "")

        if not text_content:
            print(f"⚠️ Warning: No 'answer' content found for session {session_id}. Skipping.")
            return None

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        title = doc.add_heading(f"Squash Session Plan", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Request: \"{question}\"")
        run.italic = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(24)
        lines = text_content.split('\n')
        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
            if stripped_line.startswith('###') and stripped_line.endswith('###'):
                heading_text = stripped_line.strip('#').strip()
                p = doc.add_heading(heading_text, level=1)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            elif stripped_line.startswith('•'):
                content_text = stripped_line[1:].strip()
                p = doc.add_paragraph(content_text, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)
            elif stripped_line.startswith('(Rule:'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(stripped_line)
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            elif ':' in line and any(kw in line for kw in ['Duration', 'Session Focus']):
                parts = line.split(':', 1)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.add_run(f"{parts[0]}:").bold = True
                if len(parts) > 1:
                    p.add_run(parts[1])
            elif stripped_line.startswith('---'):
                pass
            elif "End of session" in stripped_line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                run = p.add_run(stripped_line)
                run.italic = True
                run.font.size = Pt(10)
            else:
                doc.add_paragraph(line)

        output_filename = f"{session_id}.docx"
        output_path = output_dir / output_filename
        doc.save(output_path)
        return output_path
    except Exception as e:
        session_id_for_log = "N/A"
        if isinstance(session_data, dict):
            session_id_for_log = session_data.get('session_id', 'N/A')
        print(f"❌ Error processing session {session_id_for_log}: {e}")
        return None


def main():
    """Main function to parse config and orchestrate the conversion."""
    parser = argparse.ArgumentParser(
        description="Convert session answers from a JSON file into formatted .docx files."
    )
    parser.add_argument(
        "--config", type=Path, required=True, help="Path to the YAML configuration file for paths."
    )
    args = parser.parse_args()

    try:
        with open(args.config, 'r') as f:
            config = yaml.safe_load(f)
        input_path = Path(config['paths']['input_file'])
        output_path = Path(config['paths']['output_dir'])
    except (FileNotFoundError, KeyError) as e:
        print(f"❌ Error: Could not read config file or missing key. Details: {e}")
        return

    if not input_path.is_file():
        print(f"❌ Error: Input file not found at '{input_path}'")
        return

    output_path.mkdir(parents=True, exist_ok=True)
    print(f"📂 Output directory is '{output_path}'")

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            # Load the entire file at once, as it is a standard JSON list
            all_sessions_data = json.load(f)

        if not isinstance(all_sessions_data, list):
            print(f"❌ Error: Input file '{input_path}' is not a valid JSON list of session objects.")
            return

    except json.JSONDecodeError:
        print(f"❌ Error: Could not parse '{input_path}' as JSON. Please ensure it's a valid JSON file.")
        return

    successful_conversions = 0
    # Loop through the list of session data dictionaries
    for session_data in all_sessions_data:
        if isinstance(session_data, dict):
            if create_session_docx(session_data, output_path):
                successful_conversions += 1
        else:
            print(f"⚠️ Warning: Found an item in the JSON file that is not a dictionary. Skipping.")

    print(f"\n✅ Done! Successfully created {successful_conversions} .docx files.")


if __name__ == "__main__":
    main()