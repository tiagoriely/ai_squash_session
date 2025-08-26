# convert_to_docx.py
"""
Reads a .jsonl file containing squash sessions and converts each session
into a formatted .docx file.

Usage:
    python convert_to_docx.py \
        --input data/processed/squash_dataset.jsonl \
        --output data/processed/synthetic_session/docx_sessions
"""
import argparse
import json
from pathlib import Path

# Third-party library: pip install python-docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_session_docx(session_data: dict, output_dir: Path):
    """Creates and saves a single .docx file for a given squash session."""

    try:
        # Extract metadata for the document
        session_id = session_data.get("id", "unknown_session")
        meta = session_data.get("meta", {})
        focus = meta.get("family_focus", "General")
        duration = meta.get("duration", "N/A")
        # type =
        # participants =
        # primaryShots
        # SecondaryShots
        # shotTypes
        # shotSide
        # squashLevel
        # intensity
        # Fitness


        text_content = session_data.get("contents", "")

        # --- Document Creation ---
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # 1. Add a main title
        title = doc.add_heading(f"Squash Session Plan: {focus}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. Process the pre-formatted text content line by line
        lines = text_content.split('\n')

        for line in lines:
            stripped_line = line.strip()

            # Skip empty lines, as we'll handle spacing with paragraph formats
            if not stripped_line:
                continue

            # Handle Section Headings like '### Warm-up ###'
            if stripped_line.startswith('###') and stripped_line.endswith('###'):
                heading_text = stripped_line.strip('#').strip()
                p = doc.add_heading(heading_text, level=1)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)

            # Handle bullet points like '• 7 pts: Conditioned Game...'
            elif stripped_line.startswith('•'):
                # The 'List Bullet' style adds its own bullet character
                content_text = stripped_line[1:].strip()
                p = doc.add_paragraph(content_text, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)

            # Handle indented rules like '(Rule: ...)'
            elif stripped_line.startswith('(Rule:'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(stripped_line)
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color

            # Handle metadata lines like 'Duration: ...'
            elif ':' in line and any(kw in line for kw in ['Duration', 'Session Focus']):
                parts = line.split(':', 1)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.add_run(f"{parts[0]}:").bold = True
                if len(parts) > 1:
                    p.add_run(parts[1])

            # Handle separators '----'
            elif stripped_line.startswith('---'):
                # We can skip adding this or add a section break if desired
                # For simplicity, we will ignore it.
                pass

            # Handle end of session
            elif "End of session" in stripped_line:
                p = doc.add_paragraph(stripped_line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                run = p.add_run(stripped_line)
                run.italic = True
                run.font.size = Pt(10)

            # Handle any other regular lines of text
            else:
                doc.add_paragraph(line)

        # --- Save the Document ---
        output_filename = f"{session_id}.docx"
        output_path = output_dir / output_filename
        doc.save(output_path)

        return output_path

    except Exception as e:
        print(f"❌ Error processing session {session_data.get('session_id', 'N/A')}: {e}")
        return None


def main():
    """Main function to parse arguments and orchestrate the conversion."""
    parser = argparse.ArgumentParser(
        description="Convert a .jsonl dataset of squash sessions into .docx files."
    )
    parser.add_argument(
        "-i", "--input",
        type=Path,
        required=True,
        help="Path to the input .jsonl file (e.g., data/squash_dataset.jsonl)."
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        required=True,
        help="Directory where the .docx files will be saved."
    )
    args = parser.parse_args()

    # Validate input file and create output directory
    if not args.input.is_file():
        print(f"❌ Error: Input file not found at '{args.input}'")
        return

    args.output.mkdir(parents=True, exist_ok=True)
    print(f"📂 Output directory is '{args.output}'")

    # Process the dataset
    successful_conversions = 0
    with open(args.input, 'r', encoding='utf-8') as f:
        for line in f:
            try:
                session_data = json.loads(line)
                if create_session_docx(session_data, args.output):
                    successful_conversions += 1
            except json.JSONDecodeError:
                print(f"⚠️ Warning: Could not parse a line in the JSONL file. Skipping.")

    print(f"\n✅ Done! Successfully created {successful_conversions} .docx files.")


if __name__ == "__main__":
    main()